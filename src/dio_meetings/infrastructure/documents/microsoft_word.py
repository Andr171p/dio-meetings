from typing import Union

import io

import markdown

from bs4 import BeautifulSoup
from bs4.element import PageElement, NavigableString, Tag

from docx.shared import Pt
from docx import Document as WordDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from ...core.dto import Document
from ...core.base import DocumentFactory


PARAGRAPH_LINE = 50
FONT_SIZE = 24
TABLE_STYLE = "Table Grid"

SoupElement = Union[BeautifulSoup, PageElement, Tag, NavigableString]


class MicrosoftWordFactory(DocumentFactory):
    def __init__(self) -> None:
        self.document = WordDocument()

    def create_document(self, text: str) -> Document:
        self._build_document(text)
        file_buffer = io.BytesIO()
        self.document.save("")
        file_buffer.seek(0)
        return Document.from_bytes_io(
            file_buffer=file_buffer,
            file_format="docx"
        )

    def _build_document(self, text: str) -> None:
        html = markdown.markdown(text, extensions=["tables"])
        soup = BeautifulSoup(html, "html.parser")
        for element in soup.find_all(recursive=False):
            if element.name == "h1":
                self._add_heading(element.text, level=1)
            elif element.name == "h2":
                self._add_heading(element.text, level=2)
            elif element.name == "h3":
                self._add_heading(element.text, level=3)
            elif element.name == "p":
                self._add_paragraph(element.text)
            elif element.name == "ul":
                self._add_list(element, is_ordered=False)
            elif element.name == "ol":
                self._add_list(element, is_ordered=True)
            elif element.name == "blockquote":
                self._add_quote(element.text)
            elif element.name == "hr":
                self._add_horizontal_line()
            elif element.name == "table":
                self._add_table(element)

    def _add_heading(self, text: str, level: int) -> None:
        heading = self.document.add_heading(text, level=level)
        heading.style.font.size = Pt(FONT_SIZE - level * 2)

    def _add_paragraph(self, text: str) -> None:
        paragraph = self.document.add_paragraph(text)
        paragraph.style.font.size = Pt(FONT_SIZE // 2)

    def _add_list(self, element: ..., is_ordered: bool) -> None:
        for li in element.find_all("li", recursive=False):
            if is_ordered:
                self.document.add_paragraph(li.text, style="List Number")
            else:
                self.document.add_paragraph(li.text, style="List Bullet")

    def _add_quote(self, text: str) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.add_run(f'"{text}"').italic = True

    def _add_horizontal_line(self) -> None:
        self.document.add_paragraph().add_run().add_break()
        self.document.add_paragraph("―" * PARAGRAPH_LINE).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.document.add_paragraph().add_run().add_break()

    def _add_table(self, table_element: SoupElement) -> None:
        rows = table_element.find_all("tr")
        if not rows:
            return

        columns_count = max(len(row.find_all(["th", "td"])) for row in rows)

        table = self.document.add_table(rows=len(rows), cols=columns_count)
        table.style = TABLE_STYLE

        for row_idx, row in enumerate(rows):
            cells = row.find_all(["th", "td"])
            for column_idx, cell in enumerate(cells):
                table.cell(row_idx, column_idx).text = cell.get_text(strip=True)
                if cell.name == "th":
                    paragraph = table.cell(row_idx, column_idx).paragraphs[0]
                    for run in paragraph.runs:
                        run.bold = True
